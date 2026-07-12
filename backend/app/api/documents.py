from fastapi import APIRouter, UploadFile, File, Depends, HTTPException
from sqlalchemy.orm import Session
import os
import re
from typing import List
from ..api.deps import get_db
from ..models_document import Document
from ..auth import get_current_user
from ..vector_store import get_qdrant_client
from services.embedding_service.service import get_service as get_embedding_service
from services.adaptive_retrieval.service import get_service as get_adaptive_service

router = APIRouter(prefix="/documents", tags=["documents"])

STORAGE_DIR = os.path.join(os.path.dirname(__file__), '..', '..', 'storage', 'documents')
os.makedirs(STORAGE_DIR, exist_ok=True)

ALLOWED_EXTENSIONS = {'.pdf', '.docx', '.txt', '.md', '.markdown'}


def _safe_filename(filename: str) -> str:
    base = os.path.basename(filename or 'document.txt')
    return re.sub(r"[^A-Za-z0-9._-]+", "_", base)


def _extract_text(filename: str, contents: bytes) -> str:
    extension = os.path.splitext(filename)[1].lower()
    if extension in {'.txt', '.md', '.markdown'}:
        return contents.decode('utf-8', errors='ignore')
    if extension == '.pdf':
        try:
            from pypdf import PdfReader
            import io
            reader = PdfReader(io.BytesIO(contents))
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception as exc:
            raise HTTPException(status_code=422, detail=f"Could not parse PDF: {exc}")
    if extension == '.docx':
        try:
            from docx import Document as DocxDocument
            import io
            doc = DocxDocument(io.BytesIO(contents))
            return "\n".join(paragraph.text for paragraph in doc.paragraphs)
        except Exception as exc:
            raise HTTPException(status_code=422, detail=f"Could not parse DOCX: {exc}")
    raise HTTPException(status_code=400, detail="Supported files: PDF, DOCX, TXT, Markdown")


def _chunk_text(text: str, max_chars: int = 1200, overlap: int = 160) -> List[str]:
    normalized = re.sub(r"\s+", " ", text).strip()
    if not normalized:
        return []
    chunks = []
    start = 0
    while start < len(normalized):
        end = min(start + max_chars, len(normalized))
        if end < len(normalized):
            boundary = normalized.rfind(".", start, end)
            if boundary > start + max_chars // 2:
                end = boundary + 1
        chunk = normalized[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= len(normalized):
            break
        start = max(0, end - overlap)
    return chunks


def index_document_content(doc: Document):
    chunks = _chunk_text(doc.content or "")
    if not chunks:
        return 0
    vectors = get_embedding_service().embed_texts(chunks)
    points = []
    for index, vector in enumerate(vectors):
        payload = {
            "doc_id": doc.id,
            "title": doc.title,
            "filename": doc.filename,
            "chunk_index": index,
            "text": chunks[index],
        }
        points.append({"id": doc.id * 1000000 + index, "vector": vector, "payload": payload})
    get_qdrant_client().upsert(collection_name="documents", points=points)
    get_adaptive_service().refresh()
    return len(points)


@router.post('/upload')
async def upload_document(
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
):
    extension = os.path.splitext(file.filename or '')[1].lower()
    if extension not in ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=400, detail="Supported files: PDF, DOCX, TXT, Markdown")

    contents = await file.read()
    safe_name = _safe_filename(file.filename)
    path = os.path.join(STORAGE_DIR, safe_name)
    with open(path, 'wb') as f:
        f.write(contents)

    text = _extract_text(file.filename, contents)
    if not text.strip():
        raise HTTPException(status_code=422, detail="No extractable text found in document")

    db_doc = Document(
        title=file.filename,
        filename=safe_name,
        content=text,
        uploaded_by=current_user.id,
    )
    db.add(db_doc)
    db.commit()
    db.refresh(db_doc)
    indexed_chunks = index_document_content(db_doc)

    return {
        "id": db_doc.id,
        "filename": db_doc.filename,
        "title": db_doc.title,
        "indexed_chunks": indexed_chunks,
    }


@router.get('')
def list_documents(db: Session = Depends(get_db), current_user=Depends(get_current_user)):
    docs = (
        db.query(Document)
        .filter((Document.uploaded_by == current_user.id) | (Document.uploaded_by.is_(None)))
        .order_by(Document.created_at.desc())
        .all()
    )
    return [
        {
            "id": doc.id,
            "title": doc.title,
            "filename": doc.filename,
            "created_at": doc.created_at.isoformat() if doc.created_at else None,
            "characters": len(doc.content or ""),
            "preview": (doc.content or "")[:240],
        }
        for doc in docs
    ]


@router.get('/{doc_id}')
def get_document(doc_id: int, db: Session = Depends(get_db), current_user=Depends(get_current_user)):
    doc = db.query(Document).filter(Document.id == doc_id).first()
    if not doc or (doc.uploaded_by not in {None, current_user.id}):
        raise HTTPException(status_code=404, detail="Document not found")
    return {
        "id": doc.id,
        "title": doc.title,
        "filename": doc.filename,
        "created_at": doc.created_at.isoformat() if doc.created_at else None,
        "content": doc.content,
    }
